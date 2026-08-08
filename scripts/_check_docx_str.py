# -*- coding: utf-8 -*-
"""检查方案 docx 中本次完善的关键内容"""
import os
import sys

import docx

sys.stdout.reconfigure(encoding="utf-8", errors="replace")
path = [
    f
    for f in os.listdir(".")
    if f.endswith(".docx")
    and not f.startswith("~$")
    and "reserach" not in f
    and "模板填写" not in f
][0]
d = docx.Document(path)
txt = "\n".join(p.text for p in d.paragraphs)
for t in d.tables:
    for row in t.rows:
        for c in row.cells:
            txt += "\n" + c.text

checks = [
    ("23 个工具", "23 个工具" in txt),
    ("无 19 个工具", "19 个工具" not in txt),
    ("证据链闭环", "证据链闭环（2026-08）" in txt),
    ("p118 合并引用", "合并引用" in txt),
    ("python-docx 说明", "python-docx" in txt),
    ("48 处/46 处 96%", "48 处" in txt and "46 处（96%）" in txt),
    ("MinerU 费用假设", "MinerU 云 API 按平台配额/用量计费" in txt and "无配额影响" in txt),
    ("MinerU 双引擎策略", "双引擎策略——优先 MinerU" in txt),
    ("MinerU 不可用如实披露", "MinerU 不可用" in txt and "本地 markitdown" in txt),
    ("无 parse_engine 字段声称", "parse_engine 字段" not in txt),
    ("回退实测表述", "回退机制经 mineru_test_results.json 实测" in txt),
]
for name, ok in checks:
    print(("OK  " if ok else "FAIL") + "  " + name)
print("文件:", path, "| 段落数:", len(d.paragraphs))
