# -*- coding: utf-8 -*-
"""
就地编辑《材料科学文献调研Agent_算法赛初赛方案.docx》：
在 3.3.2 说明段、5.1 目标架构选型句、5.3 第三方依赖条目补上
"拟实施依赖已列入 requirements.txt" 的显式声明（与 requirements.txt 新增段落对应）。
注意：docx 为"脚本生成＋手工修订"，不可重新生成，只能原地编辑。
运行：python scripts/_edit_docx_deps_decl.py
"""
import os
import shutil
from copy import deepcopy

import docx

os.chdir(r"D:\MMLL\4.competition\2026GOAI-3")
fname = [x for x in os.listdir(".") if x.endswith(".docx") and "reserach" not in x][0]

bak = fname + ".bak-20260808"
if not os.path.exists(bak):
    shutil.copy2(fname, bak)
    print("backup ->", bak)

doc = docx.Document(fname)


def append_suffix(p, suffix):
    last = p.runs[-1] if p.runs else None
    r = p.add_run(suffix)
    if last is not None and last._r.rPr is not None:
        r._r.insert(0, deepcopy(last._r.rPr))
    return p


edits = []
for p in doc.paragraphs:
    t = p.text
    if "以上选型以复现验证为准" in t and "对应依赖" not in t:
        append_suffix(
            p,
            "上述目标架构选型（Chroma、BGE-M3、bge-reranker-v2-m3、SQLite＋属性图）均为拟实施项，"
            "对应依赖已预先列入 requirements.txt（chromadb、sentence-transformers、FlagEmbedding），"
            "实施时再实测锁定；当前原型运行不依赖这些组件。",
        )
        edits.append("3.3.2 说明段")
    elif "目标架构选型已确定默认方案" in t and "对应依赖" not in t:
        append_suffix(p, "对应依赖已列入 requirements.txt（拟实施用途，实施时实测锁定）。")
        edits.append("5.1 目标架构选型句")
    elif "第三方依赖（关键版本）" in t and "chromadb" not in t:
        append_suffix(
            p,
            "另预置拟实施依赖：chromadb==1.5.9、sentence-transformers==5.6.1、"
            "FlagEmbedding==1.4.0（复赛实施阶段启用，届时实测锁定）。",
        )
        edits.append("5.3 第三方依赖条目")

doc.save(fname)
print("edited:", edits if edits else "NONE FOUND")
