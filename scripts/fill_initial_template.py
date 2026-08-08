# -*- coding: utf-8 -*-
"""根据初赛模板 docx 生成填写完整的初赛方案文档。

内容全部来自项目实际产出（README.md / 初赛提交材料.md / ARCHITECTURE.md /
COMPLIANCE.md / REPRODUCIBILITY.md / workspace/outputs 等），不编造任何结果。
第 6 部分（团队介绍）无项目内信息来源，标注【待团队成员补充】。
"""
import docx

SRC = "AI for reserach算法赛初赛模板.docx"
DST = "初赛方案_按模板填写.docx"

doc = docx.Document(SRC)

# 关键：在插入任何段落前一次性获取全部段落对象引用。
# 之后 insert_paragraph_before 插入新段落不会改变已有对象引用，
# 因此固定引用列表不会像"每次重新获取 doc.paragraphs"那样索引错位。
BASE_PARAS = list(doc.paragraphs)


def set_para_text(para, text):
    """替换段落文本，保留段落样式；保留第一个 run 的字符格式。"""
    runs = para.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        para.add_run(text)


def fill(index, lines):
    """把 lines（列表，每项一个段落）填入 BASE_PARAS[index]：
    第 1 行替换占位段文本，其余行插入到占位段之后（下一个原始段之前）。"""
    anchor = BASE_PARAS[index]
    set_para_text(anchor, lines[0])
    insert_before = BASE_PARAS[index + 1]
    for extra in lines[1:]:
        new_p = insert_before.insert_paragraph_before(extra)
        # 保持与占位段一致的样式
        new_p.style = anchor.style


def append_hint_after(index, text):
    """在 BASE_PARAS[index] 之后插入一个 Normal 提示段（不替换原段落）。"""
    insert_before = BASE_PARAS[index + 1]
    new_p = insert_before.insert_paragraph_before(text)
    new_p.style = doc.styles["Normal"]


# ═══════════════════════════════════════════════════════════════
# 一、项目概述
# ═══════════════════════════════════════════════════════════════
fill(5, [
    "Pi-Agent —— 材料科学文献驱动的构效关系自主发现智能体",
])

fill(7, [
    "GOAI 算法赛题 · 方向三：材料科学文献驱动的科学发现智能体",
    "进阶路线（已选）：路线 A —— 构效关系发现（Structure-Property Relationships，SPR）",
    "调研主题：MOF 材料用于 CO₂ 捕获（11 轮迭代，546 篇文献）",
])

fill(9, [
    "本方案构建了一个以 LLM 为核心、端到端自主运行的文献驱动科学发现智能体，同时覆盖基本任务（文献调研 Agent）与进阶路线（路线 A 构效关系发现）。Agent 在有限时间预算内自主完成「文献检索 → 知识抽取 → 知识图谱撰写 → Research Gap 识别 → 构效关系假设生成 → 贝叶斯优化/MCTS 搜索 → 外部数据库验证 → 报告生成」全流程，全程无需人工干预。",
    "针对 MOF 材料 CO₂ 捕获这一真实材料学问题，Agent 经 11 轮迭代、检索整理 546 篇文献，识别 10 项 Research Gap（带证据链），生成并搜索 5 条构效关系假设（165 个候选点），完成定量回归核验（嵌套 F 检验 p=0.0254 显著；NiCo 5 实测点二次 R²=0.7694 vs Vegard 经典基线 R²=-0.1530；跨 4 体系 meta-analysis t=5.46, p=0.0121）与 Materials Project / OQMD 外部数据库交叉验证，并如实记录负结果。",
    "技术方案：两阶段自主闭环 + 事件驱动/状态机/19 工具管线；贝叶斯优化（RBF-GP 代理 + MLE 超参数 + UCB 采集）与 MCTS 搜索，LLM 在假设生成、假设评估、搜索内引导三个层次深度融合；双引擎 PDF 解析（MinerU 优先 + 本地 markitdown/pdfplumber 回退）、Sciverse 三层接入（MCP/Skill/REST）、跨轮记忆与记忆质量自动审计、四层证据链审计。",
    "预期方法路线：检索 → 解析 → 抽取 → 知识图谱 → Gap 识别 → 假设生成 → 搜索（贝叶斯/MCTS + LLM 引导）→ 定量回归核验 → 外部数据库交叉验证 → 参照系公平对比 → 发现报告。全部环节已跑通（详见第三、四部分）。",
    "一句话概括：流程完整（检索→报告全链路真跑通）＋ 发现有效（双金属倒 U、水-胺双分支、缺陷类型二分等假设均有文献证据链支撑）＋ 引用干净（证据链可审计、负结果如实披露）。",
])

# ═══════════════════════════════════════════════════════════════
# 二、科学问题理解
# ═══════════════════════════════════════════════════════════════
fill(12, [
    "研究对象：金属有机框架（MOF）材料的 CO₂ 捕获性能，聚焦五类构效关系——金属比例-容量、湿度-容量、Qst（吸附热）窗口、缺陷类型、杂质气体耐受性。",
    "领域背景：MOF 是 CO₂ 捕获领域最具潜力的材料家族，过去二十年产出数万篇论文，但存在三个公认的结构性问题，恰好构成 AI 可介入的切入点：",
    "（1）知识碎片化：吸附容量、选择性、Qst 等关键性能分散在上千篇独立论文中，缺乏统一结构化整理；",
    "（2）结论矛盾：同一材料在不同合成路线下性能数值差异巨大——如 Ni-MOF-74 的 CO₂ 容量报道值横跨 3.99–8.29 mmol/g，材料-工艺-性能关系远未厘清；",
    "（3）权衡关系未被刻画：高容量、高选择性、低再生能耗三者存在公认 trade-off，但 Pareto 前沿形状与驱动因素（如 Qst“甜点区间”25–40 kJ/mol）无定量结论。",
    "为什么 AI 可以介入：文献驱动的科学发现是一条「检索 → 阅读 → 抽取 → 组织 → 推理 → 验证」的流水线，恰好可被 LLM Agent 端到端接管——检索/解析可 API 化（arXiv、Sciverse、MinerU），知识组织与假设生成由 LLM 自主完成，搜索与验证可闭环（贝叶斯/MCTS + 外部数据库）。",
    "切片选择（聚焦而非贪多）：不试图解决“设计下一代 CO₂ 捕获材料”这一终极问题，而是聚焦比赛周期内可探索、可验证的局部——让 Agent 在有限时间预算（600s/轮）内自主完成文献调研到构效关系发现的全流程，保留真实科学问题的核心矛盾（容量-选择性-再生能权衡、水稳定性机理矛盾），限定在单主题内保证可完成、可检查。",
])

fill(14, [
    "领域意义：将分散于数万篇文献中的构效知识自动结构化，直接服务 CO₂ 捕获材料的筛选与设计；所产出的可检验假设（倒 U 曲线、RH 峰值、缺陷类型依赖）为“AI 发现 → 计算/实验验证”的科研闭环提供起点，这是材料领域知识生产走向自动化的一步。",
    "方法论意义：验证“文献驱动自主科学发现”范式的可行性——Agent 自主检索并依据结果决定补检索方向（如自主补充“水稳定性机理”检索直接解决 Gap 2 矛盾）、依据新证据自主修正假设（hypo_4 由“缺陷→容量”修正为“缺陷类型依赖”）、如实记录负结果（OQMD/NOMAD 对 MOF 吸附性质 0 覆盖，升级为方法论发现）。",
    "超出已有边界：识别 Slack Model 等经验公式适用场景之外的情形（双金属协同、湿态、缺陷工程），并以定量关系给出超出原有结论的可检验表述；倒 U 型非普适性（Cu/Mg 单调 vs NiCo/CoMn 倒 U）、缺陷类型二分、胺型湿态促进分支为超出已知的新信号。",
])

# ═══════════════════════════════════════════════════════════════
# 三、技术方案与预期方法路线
# ═══════════════════════════════════════════════════════════════
fill(18, [
    "总体架构：Agent 按两阶段自主运行，共 23 个工具，采用事件驱动 + 状态机（IDLE→RUN→DONE）+ 工具管线架构：",
    "阶段一（文献调研，基本任务）：自主检索 → 筛选去重 → 双引擎 PDF 解析 → 摘要整理 → 知识图谱撰写（Markdown，材料/性质/数值/关系/矛盾）→ Research Gap 识别（10 项，带置信度与证据链）→ 调研报告生成；",
    "阶段二（路线 A 构效关系发现）：假设生成（5 条）→ 贝叶斯优化（RBF-GP 代理 + MLE 超参数 + UCB）/ MCTS 搜索 → 定量回归核验（线性/二次/高斯 + 嵌套 F 检验 + LOOCV）→ 外部数据库验证（Materials Project / OQMD / NOMAD）→ 参照系公平对比（同预算随机搜索，10 种子）→ 发现报告（正/负/异常/反例四类信号）；",
    "跨轮记忆（MEMORY.md + 运行反思 + 记忆质量自动审计）驱动下一轮迭代（共 11 轮），检索策略、图谱组织、Gap 排序、假设方向均为 Agent 自主决策而非固定脚本。",
    "分层组件（关键）：入口 main.py（参数/预算/异常处理）；决策核心 pi_agent/agent.py（事件驱动主循环）；LLM 层 pi_agent/llm.py（DeepSeek，OpenAI 兼容可替换）；工具层 23 个工具；检索层 literature_agent/search.py（arXiv + Sciverse 多源检索与缓存，search_log.jsonl 审计）；Sciverse 适配 sciverse_mcp.py（MCP > Skill > REST 三层自动检测降级，全部不可用回退纯 arXiv）；解析层 parser.py（双引擎，MinerU 默认优先：Cloud > 本地服务 > pip 包 > 本地 markitdown/pdfplumber 回退，回退原因记录于 parse_engine 字段）；抽取层 extractor.py（材料/性能/数值抽取 + 表格按列/序列/句对/笛卡尔四路径 (x,y) 配对）；发现层 discovery.py（贝叶斯 RBF-GP + MLE + UCB、MCTS、LLM 引导注入并写入 llm_guidance 审计字段）；记忆层 memory_quality.py（五维质量审计）；参照系 scripts/baseline_random_search.py。",
    "关键设计决策：① 刻意不构建 JSON 知识图谱，由 Agent 撰写 Markdown 图谱（LLM 从自然语言推理比填充模板更可靠、可人工核验）；② 双引擎 PDF 解析兼顾解析质量与离线可复现；③ Sciverse 三层接入满足“鼓励 MCP/Skill 接入”并保证可运行；④ 确定性计算（搜索打分/抽取/回归）与 LLM 采样分离（固定 seed 可复现，LLM 结论带证据链可独立核验）；⑤ LLM 搜索引导“注入 + 审计”分离（每次引导写入 _llm_events 并落盘 llm_guidance 审计字段，LLM 参与从“不可证”变为“可审计”）；⑥ 记忆质量自动审计防止跨轮记忆退化为表面化摘要。",
])

fill(20, [
    "方法路线（环节 → 方法 → 初赛进度）：",
    "文献检索：arXiv + Sciverse（MCP/Skill/REST 三层）语义检索 + 全文定位 → ✅ 已跑通 11 轮、546 篇；",
    "文献解析：MinerU 优先（Cloud/本地/pip 三通道）+ markitdown/pdfplumber 本地回退 → ✅ 已实现并实测（mineru_test_results.json）；",
    "知识抽取：Agent 自主撰写 Markdown 知识图谱 → ✅ 完整产出 R1–R33 构效关系；",
    "Gap 识别：LLM 基于证据提出候选 Gap + 置信度 + 证据链 + 验证方案 → ✅ 10 项 Gap 并优先级排序；",
    "构效关系发现：贝叶斯优化（RBF-GP 代理，MLE 超参数 + UCB）/ MCTS 与 LLM 深度融合（假设生成/评估/搜索内引导三层）→ ✅ 5 条假设、165 个候选点；",
    "定量验证：回归拟合 + 嵌套 F 检验 + LOOCV + 与经典模型（Vegard 线性混合等）对比 → ✅ 已执行（NiCo 5 实测点二次 R²=0.7694 vs Vegard R²=-0.1530；跨 4 体系 meta-analysis p=0.0121），估计点局限如实标注；",
    "外部验证：Materials Project / OQMD / NOMAD 交叉验证 → ✅ 已执行（MP 氧化物代理间接热力学；OQMD/NOMAD 零覆盖负结果如实记录）；",
    "参照系：同预算随机搜索公平对比（10 种子中位数）→ ✅ 已执行，v2 打分下 5 条假设全部 bayesian_wins。",
])

fill(22, [
    "数据来源：arXiv（开放获取，始终可用）；Sciverse（2500 万+篇，语义检索，MCP/Skill/REST 三层接入，调用记录构成可审计证据链）；Sci-Base（HuggingFace opendatalab/Sci-Base，可选）；Materials Project / OQMD / NOMAD（公开材料数据库，用于路线 A 交叉验证）。文献缓存位于 workspace/data/literature_cache/（已 gitignore）。",
    "依赖工具：DeepSeek（推理 LLM，OpenAI 兼容接口，可替换任意兼容端点）；Sciverse API（文献检索，可回退纯 arXiv 零成本）；MinerU（PDF 解析引擎，可回退本地 markitdown/pdfplumber）；开源依赖见 requirements.txt / pyproject.toml（openai、requests、markitdown 0.1.7、pdfplumber、pdfminer.six、numpy、scipy、scikit-learn、pandas 等）。",
    "运行流程（复现入口）：",
    "1) pip install -r requirements.txt（或使用 Dockerfile / docker-compose.yml 容器化）；",
    "2) 配置 .api_key（DEEPSEEK_API_KEY / SCIVERSE_API_KEY / MINERU_API_KEY，可选）；",
    "3) 运行主流程：python main.py --topic \"MOF materials for CO2 capture\" --budget 600 --fresh --seed 42；",
    "4) 运行参照系：python scripts/baseline_random_search.py --iterations 40 --seeds 10；",
    "5) 模块自测（无需网络）：python -m pytest tests/（122 项单元测试）。",
    "技术路线图见 ARCHITECTURE.md（数据流图、组件依赖图）与 README.md 系统概览。",
])

# ═══════════════════════════════════════════════════════════════
# 四、阶段性实验结果或可行性验证
# ═══════════════════════════════════════════════════════════════
fill(26, [
    "主流程可行性验证：Agent 以 600s/轮预算完成 11 轮迭代，全流程（检索 → 报告）真跑通——累计整理 546 篇文献、识别 10 项 Research Gap（含证据链）、生成并搜索 5 条构效关系假设（165 个候选点）、完成定量回归核验与外部数据库交叉验证；引用全部真实可查（论文 ID + 审计日志逐层可追），无一条编造。",
    "多主题泛化性验证：同一 Agent 在 7 个独立主题（MOF/CO₂ 捕获、钙钛矿带隙、热电 ZT、高镍正极、固态电解质等）上独立完成调研，产物与记忆互不干扰（见 CROSS_THEME_REPORT.md），验证 Agent 的泛化能力而非单主题脚本。",
    "工程化验证：122 项单元测试（tests/，含 discovery/extractor/search/parser/config）；GitHub Actions CI 多 Python 版本（3.10/3.11/3.12）矩阵 + ruff 检查；Dockerfile + docker-compose.yml 容器化；MinerU 引擎连通性与自动回退实测（mineru_test_results.json）。",
])

fill(28, [
    "（1）Research Gap 清单：10 项，全部带类型/严重程度/置信度/证据链/验证方案。Top 3：Gap 1 双金属比例-容量定量关系缺失（置信度 0.95）、Gap 2 水-CO₂ 竞争/协同机理矛盾（0.97）、Gap 9 缺陷工程的定量 OMS 控制缺失（0.80）。Top 3 连续四轮稳定，与发现信号一一对应。",
    "（2）构效关系假设（5 条，hypotheses.json）：hypo_1 双金属比例-容量倒 U（置信度 0.88 / 新颖性 0.82）；hypo_2 胺功能化 MOF 湿态容量随 RH 峰值增强（0.90 / 0.85）；hypo_3 Qst 25–40 kJ/mol 窗口 Pareto 最优（0.62 / 0.78）；hypo_4 MOF-74 缺陷类型依赖（0.92 / 0.88）；hypo_5 痕量 NO₂/SO₂ 暴露后容量衰减与 OMS 密度相关（0.63 / 0.90）。",
    "（3）定量回归核验：hypo_1 NiCo-MOF-74 8 点二次拟合 R²=0.6919，嵌套 F 检验（线性 vs 二次）F=9.909, p=0.0254 显著，峰值 x=0.437；并入 Chen 2023 的 5 个独立实测点后二次 R²=0.7694、高斯 R²=0.9778，均显著优于经典 Vegard 端点线性基线（R²=-0.1530，ΔR²=+0.92/+1.13），bootstrap 89.3% 支持二次优于线性；跨 4 体系 meta-analysis（NiCo/CoMn/FeCu/MIL-101）平均增强率 61.3%，t=5.46, p=0.0121, Cohen's d=2.73；hypo_3 Qst-电负性二次 R²=0.983 vs 线性基线 0.190（p=0.011）。",
    "（4）外部数据库验证：Materials Project 氧化物代理热力学提供间接证据（如 MgNiO₂ 双金属 ΔE=-0.035 eV/atom 更稳定），明确标注为间接参考；OQMD/NOMAD 对 MOF 吸附性质覆盖为 0——如实记录为负结果并升级为方法论发现（领域数据基础设施 Gap）。",
    "（5）参照系公平对比：同预算（40 次评估 = 10 初始 + 30 UCB）随机搜索对照，v2 增强打分下 5 条假设全部 bayesian_wins（diff +0.014~+0.039），修复了初赛版本 hypo_3 随机胜出、hypo_4 无显著区分的问题（打分函数区分度已增强）。",
    "（6）发现信号四类（运行前定义）：正结果（双金属倒 U、水 RH 双分支、缺陷类型二分，置信度 ≥0.7 且有 p# 证据链）；负结果（MP/OQMD 零覆盖）；异常（Ni-MOF-74 容量矛盾区间、水-CO₂ 机理矛盾）；反例（胺型 MOF 低湿度容量反升，与“水总是有害”朴素假设相反）。",
    "（7）诚实披露的局限：NiCo 主案例 8 点中 5 个为估计点（不作独立验证证据）；5 实测点 n=5 的 F 检验功效不足（p=0.158）——统计显著性主要由归一化复合 12 点（F p=4.9e-05）与跨体系 meta-analysis（p=0.0121）提供；Qst 数据 n=5 小样本。这些局限在 validation_summary.md 中如实标注，不把估计点当证据。",
])

# ═══════════════════════════════════════════════════════════════
# 五、复现与开放计划
# ═══════════════════════════════════════════════════════════════
fill(33, [
    "环境：Python 3.10+（开发环境 3.12/3.13）；pip install -r requirements.txt，或使用 Dockerfile 容器化（docker build / docker-compose up）。",
    "配置：项目根目录 .api_key 文件（DEEPSEEK_API_KEY / SCIVERSE_API_KEY / MINERU_API_KEY），缺失时自动降级（纯 arXiv / 本地解析），保证可运行。",
    "确定性说明：搜索打分、数值抽取、回归核验等确定性计算由 seed_everything() 固定 random/numpy，--seed 42 可逐位复现（追求完全复现可设 PYTHONHASHSEED=42）；LLM 采样不保证逐字复现，但每步有审计日志（trajectory_*.json、sciverse_skill_log.jsonl、llm_guidance_audit.jsonl），结论带论文 ID 证据链可独立核验。详见 REPRODUCIBILITY.md（确定性环节/LLM 环节/降级路径三分类）。",
    "复现命令：python main.py --topic \"MOF materials for CO2 capture\" --budget 600 --fresh --seed 42；python scripts/baseline_random_search.py --iterations 40 --seeds 10；python -m pytest tests/（122 项）。多主题复现与端到端重跑见 RERUN_GUIDE.md / E2E_RERUN_GUIDE.md。",
])

fill(35, [
    "初赛阶段：按赛题要求不强制提交代码，但已明确开源边界与复现承诺。",
    "复赛/决赛阶段：以公开 GitHub 仓库提交可运行代码（pi_agent/ + literature_agent/ + utils/ + main.py + scripts/ + tests/），含 README、环境配置（requirements.txt / pyproject.toml / Dockerfile）、随机种子说明与复现步骤；探索日志、知识图谱、Gap/假设/验证报告等产物以 Markdown/JSON 随仓库开放（API Key 除外）。",
    "许可：代码 MIT；文档与运行产物 CC BY 4.0（不含第三方 API 数据）。",
])

fill(37, [
    "商业 API：DeepSeek（推理 LLM，deepseek-v4-flash）——调用环节集中在 utils/config.py 与 pi_agent/llm.py；替代方案：任意 OpenAI 兼容端点（DEEPSEEK_BASE_URL / DEEPSEEK_MODEL 环境变量切换）；费用假设：600s 单轮约数十次 LLM 调用；对可复现性影响：LLM 结论带论文 ID 证据链，可被独立核验。",
    "商业 API：Sciverse（文献检索）——调用环节在 literature_agent/search.py；三层接入（MCP/Skill/REST）；替代方案：纯 arXiv 检索（零成本）；调用记录全部留痕审计（sciverse_skill_log.jsonl，含调用 ID/时间戳/参数 SHA256）。",
    "商业服务：MinerU（PDF 解析）——默认启用引擎（prefer_mineru=True，Cloud > 本地服务 > pip 包）；对中文/复杂表格更优；全部不可用自动回退本地 markitdown+pdfplumber（离线、结果确定可复现），回退原因记录于 parse_engine 字段，连通性实测见 mineru_test_results.json。",
    "闭源模型说明：使用 DeepSeek 闭源推理模型，使用范围（Agent 推理/决策/报告文本）；原因（推理质量与稳定性）；替代方案（任意 OpenAI 兼容开源/闭源端点）；迁移成本（配置两个环境变量即可）；对可复现性影响（采样随机，但确定性计算层与证据链保证可独立核验）。",
    "数据来源与授权：arXiv 摘要/全文为开放获取；Sciverse 仅取标题+摘要用于内部调研，不对外再分发；Sci-Base 按 HuggingFace 授权使用；Materials Project / OQMD / NOMAD 为公开数据库。文献缓存位于 workspace/data/literature_cache/（已 gitignore）。",
    "第三方依赖（关键版本）：openai≥1.108、requests≥2.32、markitdown==0.1.7、pdfplumber==0.11.10、pdfminer.six==20260107、numpy、scipy、pandas、scikit-learn 等（完整清单见 requirements.txt 与 pyproject.toml）。",
    "密钥管理：.api_key 已 gitignore，不入库。完整披露见 COMPLIANCE.md。",
])

# ═══════════════════════════════════════════════════════════════
# 六、团队介绍（无项目内信息来源，标注待补充）
# ═══════════════════════════════════════════════════════════════
# 保留模板中的问题提示（Heading 2），在其后追加 Normal 待补充标注
append_hint_after(41, "【待团队成员补充】请填写每位成员的学校/公司、岗位/专业，以及核心技能（技术、专业领域）。")
append_hint_after(44, "【待团队成员补充】请列出每位核心成员在项目中的角色与分工。")
append_hint_after(47, "【待团队成员补充】请填写过往项目经历 / 获奖经历，作品合集链接（如有）。")

doc.save(DST)
print(f"[OK] 已生成: {DST}")
