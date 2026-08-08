# -*- coding: utf-8 -*-
"""检查 docx 中 546/543/46 的口径表述一致性"""
import os
import sys

import docx

sys.stdout.reconfigure(encoding="utf-8", errors="replace")
path = [
    f
    for f in os.listdir(".")
    if f.endswith(".docx")
    and not f.startswith("~$")
    and "reserach" not in f
    and "模板填写" not in f
][0]
d = docx.Document(path)
print("段落数:", len(d.paragraphs), "| 表数:", len(d.tables))
print("==== 含 546 的段落 ====")
for p in d.paragraphs:
    if "546" in p.text:
        print(f"- {p.text}")
print("==== 含 543 的段落 ====")
for p in d.paragraphs:
    if "543" in p.text:
        print(f"- {p.text}")
