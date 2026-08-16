# -*- coding: utf-8 -*-
"""patch_docx_llm_guidance.py — 方案说明文档.docx 的 LLM 引导表述补注（W-6a P2-1）

背景：方案文档技术路线描述「搜索过程中由 LLM 引导（记录在 llm_guidance 审计字段）」，
未区分主案例（事后回填审计、未影响采样）与 mof_e2e_v4（真实端到端）。文档另一处已
明确「mof_e2e_v4 是唯一带 llm_guidance 端到端审计证据的主题」，本脚本在该句末尾追加
括号注记，与 REPRODUCIBILITY.md §3.1 口径完全一致（零虚构，仅补披露）。

用法：
  python scripts/patch_docx_llm_guidance.py            # dry-run：定位目标段
  python scripts/patch_docx_llm_guidance.py --apply    # 写回（先备份 .bak-llmguid-YYYYMMDD）
"""
import glob
import os
import shutil
import sys
from datetime import date

import docx

ROOT = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..')
SRC = glob.glob(os.path.join(ROOT, '方案说明文档.docx'))[0]

NOTE = ('（注：主案例的 llm_guidance 为事后回填审计、未影响采样；'
        '仅 mof_e2e_v4 主题为真实端到端 LLM 引导，详见 docs/REPRODUCIBILITY.md §3.1）')
NEEDLE = '搜索过程中由 LLM 引导'


def main():
    dry = '--apply' not in sys.argv
    d = docx.Document(SRC)
    target = None
    for p in d.paragraphs:
        if NEEDLE in p.text:
            target = p
            break
    if target is None:
        print('未找到目标段（可能已修改或文本拆分），不处理。')
        return 1
    print('目标段:', target.text[:120])
    if NOTE in target.text:
        print('注记已存在，跳过。')
        return 0
    if dry:
        print('（dry-run，加 --apply 写入注记）')
        return 0
    bak = os.path.join(ROOT, f'方案说明文档.docx.bak-llmguid-{date.today():%Y%m%d}')
    if not os.path.exists(bak):
        shutil.copy2(SRC, bak)
        print('备份:', bak)
    target.add_run(NOTE)
    d.save(SRC)
    print('已写回注记。')
    return 0


if __name__ == '__main__':
    sys.exit(main())
