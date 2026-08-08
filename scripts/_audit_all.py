# -*- coding: utf-8 -*-
"""提交前全量核验：方案文档 vs 仓库实际文件"""
import glob
import json
import os
import re
import sys

import docx

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

DOCX = [
    f
    for f in os.listdir(".")
    if f.endswith(".docx")
    and not f.startswith("~$")
    and "reserach" not in f
    and "模板填写" not in f
][0]
d = docx.Document(DOCX)


def doc_text():
    parts = []
    for p in d.paragraphs:
        parts.append(p.text)
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                parts.append(c.text)
    return "\n".join(parts)


def read(p):
    return open(p, encoding="utf-8").read()


DT = doc_text()
results = []


def check(name, cond, detail=""):
    results.append((name, bool(cond), detail))
    print(("OK  " if cond else "FAIL") + f"  {name}" + (f"  [{detail}]" if detail else ""))


# ---- Gap 置信度一致性 ----
g = read(r"workspace/outputs/literature_survey/gap_report.md")
check("gap_report Gap1 段落=0.95", "0.95（↑ 自 0.90）" in g)
check("gap_report 优先级 Gap2=0.97", "0.97（↑0.95）" in g)
check("gap_report 优先级 Gap9=0.80", "0.80（↑0.78）" in g)
for kw, val in [
    ("双金属比例-容量定量标度", "0.95"),
    ("水-CO₂ 竞争/协同机理矛盾", "0.97"),
    ("缺陷工程定量 OMS 控制缺失", "0.80"),
]:
    i = DT.find(kw)
    check(f"方案 Gap表[{kw[:6]}]={val}", i >= 0 and val in DT[i : i + 80], DT[i : i + 50] if i >= 0 else "")

# ---- P75 证据数（应已删除） ----
check("方案已删除 3–29 条", "3–29 条" not in DT)
check("方案已删除 中位数 7", "中位数 7" not in DT)
check("方案保留置信度区间", "0.70–0.97" in DT and "均值 0.82" in DT and "100%" in DT)

# ---- P57 ----
check("README 546 篇", "546" in read("README.md"))
s4 = read(r"workspace/outputs/mof_e2e_v4/literature_survey/survey_report.md")
check("e2e_v4 149(109+40)", "149 篇检索论文（109 历史复用 + 40 本轮新增" in s4)
check("e2e_v4 Gap 1-12", "Gap 1-12" in s4 or "Gap 1–12" in s4)
check("方案含 546/149/12 项 Gap", "546" in DT and ("109 历史复用＋40 本轮新增" in DT or "109 历史复用 + 40 本轮新增" in DT))

# ---- P59 ----
ct = read("CROSS_THEME_REPORT.md")
check("CROSS 813+", "813+" in ct)
check("CROSS 51 Gap", "51" in ct)
n6 = len(re.findall(r"^### 6\.\d", ct, re.M))
check("CROSS §6 五个连接", n6 >= 5, str(n6))
check("方案含 813+/51/5 连接", "813+" in DT and "51 项 Gap" in DT and "5 个跨领域连接" in DT)

# ---- P60 ----
hj = json.load(open(r"workspace/outputs/literature_survey/discovery/hypotheses.json", encoding="utf-8"))
qs = [
    len((h.get("prior_art_verification") or {}).get("queries_executed", []))
    for h in hj
]
novs = sorted(round(h.get("novelty_score", 0), 2) for h in hj)
check("主案例 5 假设×3 查重=15", len(hj) == 5 and sum(qs) == 15, f"queries={qs}")
check("新颖性集合一致", novs == [0.78, 0.82, 0.85, 0.88, 0.90], f"{novs}")
check("方案含 15 次查重/0.78–0.90", "15 次" in DT and "0.78–0.90" in DT)

# ---- P61 ----
sh = json.load(open(r"workspace/outputs/mof_e2e_v4/literature_survey/discovery/search_h0.json", encoding="utf-8"))
g0 = sh.get("llm_guidance", {})
quote = "应剪枝温度变量，固定273K，并扩展组分覆盖两端以验证倒U形状"
evtxt = " | ".join(str(e.get("suggestion", "")) for e in g0.get("events", []))
check("llm_guidance n_events=5", g0.get("n_events") == 5, f"n={g0.get('n_events')}")
check("引语为逐字原文", quote in evtxt)
check("方案引语逐字", quote in DT)

# ---- P65 ----
qp = json.load(open(r"workspace/outputs/literature_survey/discovery/quantitative_pairs.json", encoding="utf-8"))
n24 = len(qp) if isinstance(qp, list) else len(qp.get("pairs", []))
check("quantitative_pairs=24", n24 == 24, str(n24))
check("方案含 150 篇/补标 26/参考 Gap 10", "150 篇" in DT and "补标 26 条" in DT and "参考 Gap 10 项" in DT)

# ---- P74 ----
log = read(r"workspace/logs/sciverse_skill_log.jsonl")
check("sciverse 日志 149 行", len([x for x in log.split("\n") if x.strip()]) == 149)
check("方案含 149 条", "149 条" in DT)

# ---- P76 ----
main_survey = read(r"workspace/outputs/literature_survey/survey_report.md")
check("survey_report 结构要素", "6 大类" in main_survey)
check("方案含 R1–R33/20 篇", "R1–R33" in DT and "20 篇" in DT)

# ---- P78-84 定量数字 ----
for k in ["0.7694", "0.9778", "-0.1530", "0.0254", "9.909", "0.0121", "5.46", "61.3", "0.983", "0.190", "0.011", "89.3", "4.9e-05"]:
    check(f"定量数字 {k}", k in DT)

# ---- 改写段 ----
for k in ["标准不是模型说了什么", "评分点在于", "团队拍板的默认方案", "这三条不是概念包装", "复用设计是我们认可的", "连用均值预测都不如", "关键设计选择与理由"]:
    check(f"改写段 {k[:8]}", k in DT)

# ---- p# 证据链状态（仓库级提示） ----
gaps_p = len(re.findall(r"\bp\d+\b", g))
ps = read(r"workspace/outputs/literature_survey/paper_summaries.md")
check("paper_summaries 不含 p#", not re.search(r"\bp\d+\b", ps))
print(f"INFO gap_report 中 p# 标识 {gaps_p} 处；paper_summaries 使用 r# ID（{len(re.findall(r'r\d+s\d+_\w+', ps))} 条）")

print()
print("FAIL 数:", sum(1 for _, ok, _ in results if not ok), "/", len(results))
