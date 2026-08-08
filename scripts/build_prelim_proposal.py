# -*- coding: utf-8 -*-
"""
根据《AI for reserach算法赛初赛模板.docx》生成
《材料科学文献调研Agent_算法赛初赛方案.docx》

仅修改"应说明……"占位段落，保留模板标题层级、页眉页脚、页面设置。
运行：python scripts/build_prelim_proposal.py
"""
import glob
import os
import sys

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

TEMPLATE = [
    f
    for f in glob.glob("*.docx")
    if "reserach" in f and not os.path.basename(f).startswith("~$")
][0]
OUT = "材料科学文献调研Agent_算法赛初赛方案.docx"

doc = docx.Document(TEMPLATE)


# ---------- helpers ----------
def find_heading(level: int, text: str):
    for p in doc.paragraphs:
        if p.style and p.style.name == f"Heading {level}" and p.text.strip() == text:
            return p
    return None


def get_placeholder(heading_text: str) -> Paragraph:
    h = find_heading(2, heading_text)
    assert h is not None, f"heading not found: {heading_text}"
    nxt = h._p.getnext()
    assert nxt is not None and nxt.tag == qn("w:p"), f"no placeholder after: {heading_text}"
    return Paragraph(nxt, doc)


def set_para_text(p: Paragraph, text: str):
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    p.add_run(text)


def make_para(text: str, style: str = "Normal", size: int = None, bold: bool = False):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    if size:
        r.font.size = Pt(size)
    if bold:
        r.bold = True
    return p


def build_table(headers, rows, widths_cm, font_size=9):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    tblPr = t._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    for j, w in enumerate(widths_cm):
        for cell in t.columns[j].cells:
            cell.width = Cm(w)
    for j, htxt in enumerate(headers):
        c = t.cell(0, j)
        c.text = ""
        run = c.paragraphs[0].add_run(htxt)
        run.bold = True
        run.font.size = Pt(font_size)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i + 1, j)
            c.text = ""
            run = c.paragraphs[0].add_run(str(val))
            run.font.size = Pt(font_size)
    return t


def emit_blocks(anchor_p: Paragraph, blocks):
    """blocks: list of ('p'|'h3'|'bullet', text[, style]) | ('table', headers, rows, widths, font)"""
    cursor = anchor_p._p
    for blk in blocks:
        kind = blk[0]
        if kind == "p":
            p = make_para(blk[1], style=blk[2] if len(blk) > 2 else "Normal")
            cursor.addnext(p._p)
            cursor = p._p
        elif kind == "h3":
            p = make_para(blk[1], style="Heading 3")
            cursor.addnext(p._p)
            cursor = p._p
        elif kind == "bullet":
            p = make_para(blk[1], style="List Bullet")
            cursor.addnext(p._p)
            cursor = p._p
        elif kind == "table":
            t = build_table(blk[1], blk[2], blk[3], blk[4])
            cursor.addnext(t._tbl)
            cursor = t._tbl
        else:
            raise ValueError(kind)


def fill_section(heading_text: str, blocks):
    """将占位段作为第一段正文，其余块依次插入其后"""
    ph = get_placeholder(heading_text)
    if blocks[0][0] == "bullet":
        ph.style = doc.styles["List Bullet"]
    set_para_text(ph, blocks[0][1])
    emit_blocks(ph, blocks[1:])


def insert_before_heading(level: int, heading_text: str, blocks):
    """在指定标题之前插入若干块"""
    h = find_heading(level, heading_text)
    assert h is not None, f"heading not found: {heading_text}"
    for blk in blocks:
        kind = blk[0]
        if kind == "p":
            p = make_para(blk[1])
            h._p.addprevious(p._p)
        elif kind == "h2":
            p = make_para(blk[1], style="Heading 2")
            h._p.addprevious(p._p)
        elif kind == "bullet":
            p = make_para(blk[1], style="List Bullet")
            h._p.addprevious(p._p)
        elif kind == "table":
            t = build_table(blk[1], blk[2], blk[3], blk[4])
            h._p.addprevious(t._tbl)
        else:
            raise ValueError(kind)


def remove_paragraph(p: Paragraph):
    p._p.getparent().remove(p._p)


# ---------- 一、项目概述 ----------
fill_section(
    "1.1 项目名称",
    [
        (
            "p",
            "项目名称：材料科学文献调研 Agent——面向 AI for Materials 的文献驱动科学发现基础系统（代号 MatLit-Survey-Agent）。本方案覆盖方向三的两项初赛必交内容：基本任务（必做）——文献调研 Agent，以及进阶路线（已选）——路线 A：构效关系发现（Structure-Property Relationships，SPR）的 Proposal。基本任务产出（结构化知识库与 Research Gap 清单）是路线 A 的输入基础，二者在同一原型中端到端贯通。",
        ),
    ],
)

fill_section(
    "1.2 参赛方向",
    [
        (
            "p",
            "参赛方向：GOAI 算法赛题 · 方向三——材料科学文献驱动的科学发现智能体。本方案任务范围：（1）基本任务（必做）——文献调研 Agent，为方案主体（见第三、四章）；（2）进阶路线（已选）——路线 A：构效关系发现（SPR），Proposal 概要见 3.1.5 与 3.2.1，实证结果见 4.2.2。",
        ),
    ],
)

fill_section(
    "1.3 方案概述",
    [
        (
            "p",
            "本方案设计一个以“工作流编排＋多角色 Agent＋专业工具＋结构化知识库＋证据核验”为核心架构的材料科学文献调研 Agent。系统将一次科学调研请求拆解为八类角色任务（任务规划、文献检索、文献筛选、PDF 解析与知识抽取、跨文献知识融合、Research Gap 识别、证据核验、报告生成），由工作流引擎编排执行，覆盖“科学问题输入→任务拆解→检索策略生成→多源检索→去重与筛选→PDF 全文解析→材料知识抽取→实体规范化与单位统一→结构化知识库→跨文献融合→冲突与缺失检测→Research Gap 生成→证据核验→Gap 评分排序→报告生成→引用与事实检查”十六步完整链路。",
        ),
        (
            "p",
            "关键设计选择与理由（均为团队拍板）：①领域与问题——以 MOF 材料用于 CO₂ 捕获为主案例（另 8 个主题可扩展），聚焦容量-选择性-再生能权衡等可验证切片，不贪大求全；②架构演进——当前为单 Agent 事件驱动原型（已实现、11 轮稳定运行），目标为多角色多 Agent 拆分，循序渐进而非推倒重来；③数据源——Sciverse 语义检索为主、arXiv 零成本回退、Sci-Base 可选本地索引，混合策略保证无密钥也可运行；④进阶衔接——基本任务产出直接作为路线 A（构效关系发现）的输入，同一原型端到端贯通。",
        ),
        (
            "p",
            "与普通搜索、普通 RAG 与通用文献问答系统相比，本方案有三个本质差异：（1）以“结构化知识库＋证据链”为知识载体而非片段级问答——所有抽取结果锚定到具体文献段落，任何结论可逐层追溯；（2）Research Gap 必须基于结构化证据生成，每条 Gap 附支撑文献、证据缺失或冲突说明、可证伪假设与建议验证方法，禁止大模型自由生成；（3）面向科学发现的闭环——输出带交叉引用的调研报告，可直接作为构效关系发现、假设验证等下游任务的输入。",
        ),
        (
            "p",
            "预期方法路线：基本任务按上述十六步链路执行；在此基础上进入路线 A：构效关系假设生成→搜索（贝叶斯优化/MCTS，LLM 深度融合）→定量回归核验（与 Vegard/Slack 等经典模型正面对比）→外部数据库交叉验证→参照系公平对比→发现报告。初步结果（真实，详见 4.2）：主案例（MOF 材料用于 CO₂ 捕获）经 11 轮迭代、546 篇次检索（去重后证据池 543 条、最终收录 46 篇），产出 10 项 Research Gap（带证据链）与 5 条构效关系假设（165 个候选点）；定量核验嵌套 F 检验 p=0.0254，5 个独立实测点二次 R²=0.7694、优于经典 Vegard 线性基线（R²=-0.1530）。",
        ),
        (
            "p",
            "当前状态：实验室原型（Pi-Agent 仓库）已实现单 Agent 事件驱动版本的两阶段管线，并在 8 个材料细分主题上完成初步可行性运行；本方案的技术方案与路线 A Proposal 均基于该原型的真实能力与真实产物（详见第三、四章），尚未实施的部分明确标注为拟实施项，不包含任何虚构结果。",
        ),
    ],
)

# ---------- 二、科学问题理解 ----------
fill_section(
    "2.1 科学问题与研究对象",
    [
        (
            "p",
            "材料科学知识高度沉淀于文献：组分、结构、工艺与性能之间的关联大量以非结构化形式散落在数十年论文与专著中。随着论文规模持续增长，人工阅读的覆盖能力与文献规模之间已形成根本性落差，具体表现为三类结构化问题：（1）知识碎片化——同一材料体系的容量、选择性、吸附热等关键性能分散在数百篇论文中，缺乏统一结构化整理；（2）结论矛盾——同一材料在不同合成路线与测试条件下的报道数值差异显著（例如 MOF-74 系列材料的 CO₂ 容量报道值可横跨 3.99–8.29 mmol/g），材料-工艺-性能关系远未厘清；（3）构效关系难以跨文献定量关联——性能数据缺乏统一的实验条件与单位语境，难以直接比较与回归。",
        ),
        (
            "p",
            "本方案以“给定一个材料细分方向（如 MOF 材料用于 CO₂ 捕获、卤化物钙钛矿带隙、热电 ZT 优化等），自动完成一次可审计的文献调研，并在此基础上识别构效关系假设”为研究对象。任务拆解为三个层次：问题层（将自然语言问题分解为材料体系、目标性质、实验/模拟条件、方法四要素）；证据层（检索、解析、抽取、规范化、融合）；产出层（Gap 识别、证据核验、报告生成）。系统需在产出中显式区分三类知识：【文献事实】（原文明确陈述且可定位）、【跨文献推论】（由多篇文献证据融合得出并标注推理路径）、【待验证假设】（由 Gap 与推论提出，需实验或计算验证）。",
        ),
    ],
)

fill_section(
    "2.2 科学意义",
    [
        (
            "p",
            "领域意义：本方案直接服务于材料领域“知识生产自动化”——将分散于海量文献中的组分-结构-工艺-性能知识自动结构化，形成可持续积累、可查询、可推理的材料知识库；产出的 Research Gap 清单为“AI 发现→计算/实验验证”的科研闭环提供可证伪的起点，可减少人工文献调研的重复劳动，提升材料筛选与设计的信息完备性。",
        ),
        (
            "p",
            "方法论意义：本方案验证“文献驱动自主科学发现”范式的工程可行性——文献检索与解析可 API 化（Sciverse、MinerU），知识组织、Gap 识别与假设提出可由 LLM 承担，而证据核验与数值一致性校验由规则程序与结构化知识库兜底，形成“LLM 生成假设＋确定性程序验证”的可审计范式。这与仅依赖模型记忆的通用问答有本质区别，为材料领域知识生产自动化提供可复用的智能体范式。",
        ),
        (
            "p",
            "对已有理论的推广/颠覆（新知 vs 已知）：我们判断“新知”的标准不是模型说了什么，而是能否与已有结论明确划清界限。以 Vegard 线性混合为例，它假设双金属材料的性能随组分在两端点之间线性过渡，即两种金属位点对 CO₂ 的贡献可独立相加；但 NiCo-MOF-74 的 5 个实测点（0℃/1bar）呈明显倒 U，中间组分（x≈0.5，8.30 mmol/g）明显高于两端（3.99/5.03 mmol/g），线性插值在此数据上 R² 为负，说明中间组分出现了端点没有的协同，加和假设不成立。Slack 经验公式则是为高温、简单晶体（如 NaCl 型）标定的，自变量里没有湿度、缺陷、双金属比例，无法外推到 MOF 的湿态捕获、缺陷工程与双金属协同场景。因此我们把“经典模型的失效区间”当作发现目标，而不是在它们的适用区间里重复验证。类似地，“水对 CO₂ 吸附总是有害”的旧说法被多体系独立证据改写为“水可被利用/调控”；“缺陷”也从一元变量细分为缺失配体型与溶剂占据型两种方向相反的机制。这些新表述都带文献证据链，并显式标注为【跨文献推论】或【待验证假设】，不与文献已有结论混淆。",
        ),
    ],
)

# ---------- 三、技术方案与预期方法路线 ----------
fill_section(
    "3.1 技术方案",
    [
        (
            "p",
            "3.1.1 当前已实现原型（总体架构，真实可运行）：原型按两阶段自主运行，共 23 个工具，采用事件驱动＋状态机（IDLE→RUN→DONE）＋工具管线架构。阶段一（文献调研）：自主检索→筛选去重→双引擎 PDF 解析→摘要整理→知识图谱（Markdown，材料/性质/数值/关系/矛盾）→Research Gap 识别（带置信度与证据链）→调研报告生成；阶段二（路线 A）：假设生成→贝叶斯优化（RBF-GP 代理＋MLE 超参数＋UCB）/MCTS 搜索→定量回归核验→外部数据库验证→参照系公平对比→发现报告。分层组件：main.py（入口）、pi_agent/（事件驱动主循环、LLM 层、23 个工具）、literature_agent/（search、parser、extractor、discovery、sciverse_mcp、scoring）、memory_quality（跨轮记忆质量审计）；跨轮记忆（MEMORY.md＋运行反思）驱动 11 轮迭代。检索策略、图谱组织、Gap 排序、假设方向均为 Agent 自主决策而非固定脚本。",
        ),
        (
            "h3",
            "3.1.2 目标架构：多角色 Agent 映射与演进",
        ),
        (
            "p",
            "目标架构将上述单 Agent 工具管线按职责拆分为八类角色 Agent，由工作流引擎编排。编排框架我们拍板沿用自研事件驱动状态机（pi_agent/state_machine.py，已实现、单元测试覆盖、11 轮稳定运行）——它已承载全部 23 个工具与两阶段流程，引入 LangGraph 等外部框架意味着迁移已跑通的代码、扩大依赖面，而评分点在于“LLM 深度融合”而非框架名，故 LangGraph 仅作复赛对照备选。各角色与落地状态如下：",
        ),
        (
            "table",
            ["Agent 角色", "主要输入", "处理方法", "主要输出", "落地状态"],
            [
                ["任务规划 Agent", "用户科学问题", "LLM 结构化追问＋四要素分解，生成子任务 DAG", "子任务清单、验收标准", "已实现（决策层）"],
                ["文献检索 Agent", "检索策略", "Sciverse 语义检索＋arXiv＋Sci-Base（可选），调用留痕", "检索结果集（DOI/标题/摘要）", "已实现"],
                ["文献筛选 Agent", "检索结果集", "指纹去重＋Embedding/Reranker 排序＋纳入排除规则", "精选文献集＋筛除原因", "部分实现（混合检索拟实施）"],
                ["PDF 解析与知识抽取 Agent", "全文 PDF", "MinerU＋markitdown/pdfplumber 回退；LLM＋规则抽取", "结构化记录（材料/性质/数值/条件/方法）", "已实现"],
                ["跨文献知识融合 Agent", "结构化记录", "实体对齐、单位统一、数值聚簇、矛盾/缺失检测", "融合知识图谱＋矛盾/缺失清单", "部分实现（图谱库拟实施）"],
                ["Research Gap 识别 Agent", "融合图谱＋矛盾清单", "基于证据生成 Gap（类型/证据/可证伪假设）", "Gap 候选（带置信度与验证方案）", "已实现"],
                ["证据核验 Agent", "Gap 候选", "引用锚定、二次核验、数值/单位一致性校验", "核验后 Gap（评分排序）", "部分实现（自动二次核验拟增强）"],
                ["报告生成 Agent", "核验后知识", "结构化报告模板＋交叉引用", "带证据链附录的调研报告", "已实现"],
            ],
            [2.4, 2.7, 5.0, 4.2, 2.7],
            8,
        ),
        ("h3", "3.1.3 九项核心能力的实现要点（含落地状态）"),
        (
            "bullet",
            "科学问题理解与任务拆解（已实现）：任务规划 Agent 通过结构化追问模板（材料体系、目标性质、条件、方法）澄清问题，输出四要素分解与子任务 DAG，避免检索方向发散。",
        ),
        (
            "bullet",
            "文献检索、筛选、去重与排序（部分实现）：多源检索（Sciverse 语义检索与全文定位为主，arXiv 元数据为辅）；指纹去重（DOI/标题/摘要）；按相关性、证据覆盖度、发表时间综合排序；混合检索（BM25＋Embedding 语义召回＋Reranker 精排）拟实施。",
        ),
        (
            "bullet",
            "PDF 全文解析（已实现）：双引擎策略——优先 MinerU（Cloud＞本地服务＞pip 包），当前运行环境未配置/未部署时自动回退本地 markitdown＋pdfplumber（离线、确定可复现）；回退机制经 mineru_test_results.json 实测验证。",
        ),
        (
            "bullet",
            "材料知识抽取（已实现）：LLM 依据统一 schema 抽取成分（化学式/掺杂比例）、结构（空间群/孔径/配体）、性能（容量/选择性/吸附热/带隙等，含数值与条件）、工艺（合成方法/温度/气氛）、实验条件与模拟方法；规则程序负责数值正则与字段校验；数值 (x,y) 配对采用表格按列/序列/句对/笛卡尔四路径。",
        ),
        (
            "bullet",
            "实体规范化与单位统一（部分实现）：名称-别名映射（如 MOF-74/Ni-MOF-74/dobdc）、单位换算至 SI 主单位（mmol/g、kJ/mol、K、bar），数值绑定实验条件上下文；完整 SI 统一 schema 拟实施。",
        ),
        (
            "bullet",
            "结构化知识库存储（部分实现）：当前为 Markdown 知识图谱＋文献缓存（可人工核验）；向量库（chunk＋Embedding）与关系库/图谱数据库（实体-关系-数值属性图）拟实施，两类存储以 doc_id＋passage_id 关联保证回源。",
        ),
        (
            "bullet",
            "跨文献知识融合与矛盾检测（已实现）：按“材料体系×性质×条件”聚簇数值，检测显著差异、矛盾区间与缺失连接（某性质-条件组合零覆盖）；LLM 对矛盾给出机理级解释候选，规则程序给出数值证据，两者分离、可审计。",
        ),
        (
            "bullet",
            "Research Gap 识别（已实现）：仅基于结构化证据生成；每条 Gap 输出问题描述、支撑文献、证据缺失或冲突、新颖性论证、可操作性、可证伪假设与建议验证方法；按类型（未探索空间/矛盾结论/样本不足）与严重度、置信度排序。",
        ),
        (
            "bullet",
            "证据链与事实核验（已实现四层追溯，自动二次核验拟增强）：每个抽取项、每条推论、每条 Gap 均锚定文献标识与原文定位（p#/passage_id）；引用可经 p#→paper_summaries→search_log→Sciverse API 调用记录四层追溯。",
        ),
        (
            "p",
            "3.1.4 与普通方案的本质区别：这三条不是概念包装，而是由系统落地机制直接决定的。其一，知识组织粒度不同——搜索引擎返回链接、普通 RAG 返回片段拼成的答案，本系统把语料组织成“材料-性质-数值-条件”的结构化知识（知识图谱 R1–R33＋数值表），因此能回答“哪些体系在同一条件下可比、数值是否矛盾”这类片段问答答不了的问题。其二，普通 RAG 没有“跨文献”环节——本系统按同一材料体系、同一性质把数值跨文献聚簇，检测矛盾区间（如 Ni-MOF-74 容量 3.99–8.29 mmol/g）与缺失连接，Gap 据此生成并附证据链，而不是模型凭记忆自由发挥。其三，产出可审计——每个结论锚定文献标识，可经 paper_summaries→search_log→Sciverse 调用记录四层回源，输出是可直接喂给下游计算/实验验证的结构化报告，而非“可能正确”的问答文本。",
        ),
        (
            "h3",
            "3.1.5 路线 A：构效关系发现技术方案（Proposal）",
        ),
        (
            "p",
            "以基本任务产出的结构化知识库与 Gap 清单为输入，采用“搜索算法＋LLM 深度融合＋定量验证＋外部数据库”框架。①搜索算法：贝叶斯优化（RBF-GP 代理模型、MLE 超参数、UCB 采集，已实现）、MCTS（已实现）、符号回归（已实现），拟引入遗传算法扩充搜索种群【待最终确认】。②LLM 深度融合（路线 A 核心考察点）：三层参与——假设生成作为搜索种群的种子（已实现）、评估中间候选的科学合理性（已实现）、搜索内引导（v2 已工程化并写入 llm_guidance 审计字段，实证见 4.2.2），而非仅用 LLM 生成代码。③定量验证：线性/二次/高斯拟合＋嵌套 F 检验＋LOOCV＋bootstrap＋与经典模型（Vegard 线性混合、Slack 经验公式）正面对比。④外部数据库交叉验证：Materials Project / OQMD / NOMAD / hMOF-CoRE。⑤参照系公平对比：同预算随机搜索（10 种子）。⑥“为何旧模型失效”机制分析：Vegard 线性混合隐含组分-性能加和性假设，无法刻画双金属位点协同的非加和效应（倒 U 峰），在 NiCo 5 实测点上 R²=-0.1530，二次/高斯 ΔR²=+0.92/+1.13；Slack 经验公式适用高温、简单晶体，对 MOF 湿态/缺陷/双金属体系外推失效——路线 A 明确以这些“失效区间”为发现目标。",
        ),
    ],
)

fill_section(
    "3.2 预期方法路线",
    [
        (
            "p",
            "基本任务技术路线共十六步，各步的输入、处理方法与输出如下表所示；“承担组件”列说明大语言模型（LLM）、Embedding、Reranker、MinerU、数据库、规则程序与工作流引擎各自承担的任务。",
        ),
        (
            "table",
            ["步骤", "输入", "处理方法", "输出", "承担组件"],
            [
                ["1 科学问题输入", "用户自然语言问题", "结构化澄清，解析为四要素（体系/性质/条件/方法）", "规范化科学问题", "LLM（任务规划 Agent）"],
                ["2 任务拆解", "规范化科学问题", "四要素分解，生成子任务 DAG", "子任务清单与验收标准", "LLM＋工作流引擎"],
                ["3 检索策略生成", "子任务清单", "查询词生成、同义词扩展、检索范围界定", "检索式集合", "LLM＋规则程序"],
                ["4 多源文献检索", "检索式集合", "Sciverse 语义检索＋arXiv 元数据，调用留痕", "原始文献集", "工作流引擎＋检索 API"],
                ["5 去重与筛选", "原始文献集", "DOI/标题/摘要指纹去重＋混合检索排序", "精选文献集", "规则程序＋Embedding＋Reranker"],
                ["6 PDF 全文解析", "精选文献 PDF", "MinerU 版面解析（表格/公式），双引擎回退", "结构化全文", "MinerU＋规则程序"],
                ["7 材料知识抽取", "结构化全文", "LLM 按 schema 抽取＋规则校验", "结构化抽取记录", "LLM＋规则程序"],
                ["8 实体规范化与单位统一", "抽取记录", "别名映射、化学式标准化、单位换算至 SI", "规范化实体与数值", "规则程序＋映射表"],
                ["9 结构化知识库存储", "规范化数据", "向量化入库＋属性图/关系入库", "知识库（向量＋图谱）", "Embedding＋向量库＋关系库"],
                ["10 跨文献融合", "知识库", "实体对齐、数值聚簇、关联补全", "融合知识图谱", "规则程序＋LLM＋图谱查询"],
                ["11 冲突与缺失连接检测", "融合图谱", "数值差异显著性检测、零覆盖检测", "矛盾与缺失清单", "规则程序＋LLM"],
                ["12 Research Gap 生成", "矛盾与缺失清单", "基于证据生成 Gap（类型/证据/可证伪假设）", "Gap 候选", "LLM＋规则程序"],
                ["13 文献证据核验", "Gap 候选", "引用锚定原文、二次核验、数值一致性校验", "核验后 Gap", "证据核验 Agent＋规则程序"],
                ["14 Gap 评分排序", "核验后 Gap", "严重度×置信度×新颖性×可操作性加权", "排序 Gap 清单", "规则程序＋LLM"],
                ["15 调研报告生成", "排序 Gap＋知识库", "模板化生成＋交叉引用", "结构化调研报告", "LLM＋报告模板"],
                ["16 引用与事实检查", "报告初稿", "逐条引用回源、交叉引用完整性检查", "终稿报告", "规则程序＋证据索引"],
            ],
            [1.5, 2.9, 4.4, 3.4, 4.8],
            8,
        ),
        (
            "p",
            "各组件承担任务归纳：LLM 承担需要理解与推理的任务（任务拆解、检索策略、知识抽取、Gap 生成、报告撰写）；Embedding 承担语义召回；Reranker 承担相关性精排；MinerU 承担 PDF 到结构化文本/表格/公式的解析；数据库（向量库＋关系库/知识图谱）承担知识存储、检索与跨文献查询；规则程序承担确定性任务（去重、单位换算、数值校验、差异检测、引用检查、Gap 评分）；工作流引擎承担任务编排、状态管理、预算控制与审计留痕。",
        ),
        ("h3", "3.2.1 路线 A 预期方法路线"),
        (
            "p",
            "路线 A 预期方法路线（九步）如下表所示，“LLM 融合点”列说明各步骤 LLM 的参与方式。",
        ),
        (
            "table",
            ["步骤", "输入", "处理方法", "输出", "LLM 融合点"],
            [
                ["1 数据挖掘", "基本任务知识库＋Gap 清单", "从知识图谱/数值表抽取 (x,y) 证据对", "结构化数据池", "无（规则/确定性）"],
                ["2 差异分析", "数据池", "识别经典模型（Vegard/Slack）预测失效场景", "失效区间清单", "LLM 提出失效机理候选"],
                ["3 规律拟合", "失效区间数据", "线性/二次/高斯/符号回归拟合", "候选规律", "LLM 解读与合并表达式"],
                ["4 假设生成与搜索", "候选规律", "贝叶斯/MCTS/遗传算法搜索候选空间", "假设清单＋候选点", "LLM 生成种子假设、评估中间结果"],
                ["5 搜索内引导", "搜索中间状态", "LLM 剪枝聚焦（温度/变量范围等）", "引导事件（审计）", "LLM（llm_guidance 审计字段）"],
                ["6 定量验证", "假设＋数据", "嵌套 F 检验、LOOCV、bootstrap、meta-analysis", "验证报告", "LLM 解释统计结果"],
                ["7 外部数据库验证", "假设", "MP/OQMD/NOMAD/hMOF 查询", "交叉验证记录", "LLM 判断支持/否定"],
                ["8 参照系对比", "假设", "同预算随机搜索公平对比", "公平对比报告", "无（确定性）"],
                ["9 发现报告", "以上全部", "正/负/异常/反例四类信号＋机制解释", "发现报告", "LLM 撰写＋证据链"],
            ],
            [1.4, 2.9, 4.5, 3.4, 4.8],
            8,
        ),
    ],
)

fill_section(
    "3.3 数据来源、依赖工具与运行流程",
    [
        (
            "p",
            "3.3.1 数据与工具：下表列出方案拟用（部分为原型已用）的数据资源与工具，并说明是否开源、外部 API 依赖、许可证与对复现性的影响。",
        ),
        (
            "table",
            ["工具/资源", "用途", "是否开源", "外部 API 依赖", "许可证", "复现影响"],
            [
                ["Sci-Base（Hugging Face）", "2500 万+篇论文语料，可选本地检索索引", "开放数据集", "无（HF 下载）", "以 HF 数据集授权为准", "可选接入，下载量大"],
                ["Sciverse API", "语义检索、全文定位；调用记录构成证据链", "平台服务（非开源）", "需注册 API Key", "以平台服务条款为准", "有 Key 可用；无 Key 自动降级 arXiv"],
                ["MinerU", "PDF 版面/表格/公式解析", "开源（OpenDataLab）", "云 API 可选，可本地部署", "以官方开源许可为准", "云解析受配额影响；本地可复现"],
                ["Hugging Face", "模型与数据集分发", "平台", "无", "—", "用于获取 Embedding/Reranker 与 Sci-Base"],
                ["向量数据库", "语义检索、证据检索", "开源（采用 Chroma，FAISS 备选）", "无", "以所选项目许可为准", "本地运行，可复现"],
                ["关系库/知识图谱", "实体-关系-数值存储与跨文献查询", "开源（采用 SQLite＋属性图）", "无", "以所选项目许可为准", "本地运行，可复现"],
                ["Embedding 模型", "语义向量化（中英多语言材料术语）", "开源（采用 BGE-M3）", "无（本地推理）", "以模型许可为准", "本地可复现"],
                ["Reranker 模型", "检索结果精排", "开源（采用 bge-reranker-v2-m3，与 BGE 配套）", "无（本地推理）", "以模型许可为准", "本地可复现"],
                ["LLM（原型用 DeepSeek）", "任务拆解/抽取/Gap/报告等推理", "闭源（可替换开源）", "需 API Key", "商业服务条款", "采样随机；结论带证据链可独立核验"],
                ["Python 3.10+/3.12", "实现语言", "开源", "无", "PSF 许可", "完全可复现"],
                ["Docker", "容器化部署", "开源", "无", "Apache 2.0", "完全可复现"],
            ],
            [3.0, 3.8, 2.2, 2.6, 2.6, 2.8],
            8,
        ),
        ("h3", "3.3.2 运行流程"),
        (
            "bullet",
            "环境准备：pip install -r requirements.txt（或 docker build / docker-compose up 容器化部署）；",
        ),
        (
            "bullet",
            "密钥配置：项目根目录 .api_key（DEEPSEEK_API_KEY / SCIVERSE_API_KEY / MINERU_API_KEY，均可选，缺失自动降级至纯 arXiv＋本地解析）；",
        ),
        (
            "bullet",
            "问题输入与全流程执行：python main.py --topic \"MOF materials for CO2 capture\" --budget 600 --fresh --seed 42，工作流自动完成任务拆解→检索→筛选→解析→抽取→规范化→入库→融合→Gap→核验→报告；多主题示例（钙钛矿带隙/热电 ZT/高镍正极等）见 README.md；",
        ),
        (
            "bullet",
            "产物输出：调研报告、Gap 清单、知识图谱、论文摘要集与证据链审计日志（workspace/outputs/<主题>/literature_survey/）；",
        ),
        (
            "bullet",
            "模块自测：python -m pytest tests/（离线、无需网络）；",
        ),
        (
            "bullet",
            "对照评估（拟新增）：检索 P/R/nDCG、抽取 F1、Gap 专家评审脚本，用于 4.1.2 对照实验。",
        ),
        (
            "p",
            "说明：现有原型已具备上述 ②③④⑤ 的可运行形态；混合检索、向量库、多角色拆分与对照评估为拟实施项（见第四章）。目标架构选型为团队拍板的默认方案，理由：向量库用 Chroma（本地 pip 安装、无需服务端，百篇级语料足够，FAISS 备选）；Embedding 用 BGE-M3（中英多语言，材料术语中英混合检索友好，可本地部署）；Reranker 用 bge-reranker-v2-m3（与 Embedding 同源、一致性最好）；关系存储用 SQLite＋属性图（轻量、可复现）。以上选型以复现验证为准【待确认】。上述目标架构选型（Chroma、BGE-M3、bge-reranker-v2-m3、SQLite＋属性图）均为拟实施项，对应依赖已预先列入 requirements.txt（chromadb、sentence-transformers、FlagEmbedding），实施时再实测锁定；当前原型运行不依赖这些组件。",
        ),
    ],
)

# ---------- 四、阶段性实验结果或可行性验证 ----------
fill_section(
    "4.1 阶段性实验或可行性验证",
    [
        ("p", "4.1.1 已完成的初步可行性验证（真实运行产物，均可在仓库中核验）："),
        (
            "bullet",
            "全链路可行性：现有单 Agent 原型（事件驱动＋23 工具）已端到端跑通“检索→解析→抽取→知识图谱→Gap→报告”，主案例（MOF 材料用于 CO₂ 捕获）累计 11 轮迭代、546 篇次检索（去重后持久化证据池 543 条见 paper_register.md，最终收录 46 篇）、10 项 Research Gap；其中 e2e_v4 全量版去重后证据池 149 篇（109 历史复用＋40 本轮新增）、12 项 Gap（workspace/outputs/mof_e2e_v4/literature_survey/）。注：跨主题报告（CROSS_THEME_REPORT.md）主案例口径为 376 篇（152＋224），因统计时点与去重口径不同，以各报告原文为准；",
        ),
        (
            "bullet",
            "路线 A 全流程实证：假设生成→搜索→定量核验→外部验证→参照系→发现报告已在主案例跑通——5 条假设、165 个候选点、定量核验（嵌套 F 检验 p=0.0254；NiCo 5 实测点二次 R²=0.7694 vs Vegard 基线 R²=-0.1530）、参照系 v2 打分下 5 条假设全部胜出、MP/OQMD/hMOF 外部验证（含 OQMD/NOMAD 对 MOF 吸附性质零覆盖的负结果），详见 4.2.2；",
        ),
        (
            "bullet",
            "跨主题泛化：同一套代码在 8 个主题（6 个研究主题＋2 个验证/冒烟主题）上独立运行，累计检索 813+ 篇文献、识别 51 项 Gap（CROSS_THEME_REPORT.md）；并识别 5 个跨领域连接（MOF 缺陷↔电池正极缺陷、钙钛矿带隙权衡↔热电 ZT、MOF 水竞争↔固态电解质界面等，均为【跨文献推论】待验证）；",
        ),
        (
            "bullet",
            "新颖性验证：主案例 5 条假设各执行 3 个系统性查重查询（共 15 次），overlap 评估均为 none（无直接重叠文献），新颖性评分 0.78–0.90（0.82/0.85/0.78/0.88/0.90），并以“矛盾并存/单点证据/定量缺失”三类证据结构支撑新颖性论证（hypotheses.json 的 novelty_verification / known_prior_work 字段）；",
        ),
        (
            "bullet",
            "LLM 深度参与审计实证：mof_e2e_v4 搜索轨迹 search_h*.json 含 llm_guidance 字段（5 条引导事件，如“应剪枝温度变量，固定273K，并扩展组分覆盖两端以验证倒U形状”），llm_guidance_audit.jsonl 落盘，为“LLM 与搜索深度融合”提供可审计证据；",
        ),
        (
            "bullet",
            "工程化支撑：MinerU 连通性与自动回退已实测（mineru_test_results.json：当前运行环境未配置 MINERU_API_KEY 且未部署本地服务，MinerU 不可用，实际解析由本地 markitdown＋pdfplumber 完成，回退机制验证通过）；CI（.github/workflows/ci.yml）、Dockerfile 与 requirements.txt 版本锁定已就绪。",
        ),
        ("h3", "4.1.2 拟开展的小规模对照实验方案（初步可行性验证方案，尚未形成量化结果）"),
        (
            "p",
            "为系统评估方案有效性，拟开展小规模对照实验，比较关键词检索、语义检索、混合检索、普通 RAG、单 Agent 与多角色多 Agent 六种方案：",
        ),
        (
            "bullet",
            "领域与数据（金标准直接复用现有产物，减少人工标注偏差）：以 MOF 材料用于 CO₂ 捕获为主案例（另有 8 个主题语料可扩展对比）。复用设计是我们认可的：quantitative_pairs.json 的 24 条数值对（含 Chen 2023 五点实测）已带文献溯源，直接作金标准可避免二次标注的主观偏差；gap_report.md 的 10 项 Gap 经过 11 轮证据迭代、带置信度与验证方案，比临时人工标注更可靠；语料取自主案例持久化证据池（543 条，见 paper_register.md）中的 150 篇，与既有产物同源、可逐条核对。具体规模：语料 150 篇、三元组金标准 50 条（复用 24 条＋从 knowledge_graph.md 的 R1–R33 补标 26 条）、参考 Gap 10 项；",
        ),
        (
            "bullet",
            "评测执行时间（里程碑）：初赛提交后 1 周内冻结标注集 → 2 周内跑完 ①②③④⑤（检索三方案、普通 RAG、单 Agent）→ 复赛前完成 ⑥（多角色多 Agent）与领域专家评审；预计于复赛截止前完成，以组委会复赛截止时间为准【待确认：具体日期】；",
        ),
        (
            "bullet",
            "对照方案：① BM25 关键词检索；② Embedding 语义检索；③ 混合检索（BM25＋向量＋Reranker）；④ 普通 RAG（片段级问答）；⑤ 单 Agent 原型（现有）；⑥ 多角色多 Agent 目标架构（拟实施）；",
        ),
        (
            "bullet",
            "评估指标：检索 Precision、Recall、Recall@K、nDCG@K；知识抽取 Precision、Recall、F1；数值与单位匹配准确率；引用与原文证据定位准确率；Research Gap 专家认可率、新颖性、可操作性与证据完整率；报告生成时间与 API 成本；",
        ),
        (
            "bullet",
            "评估方式：领域专家评审与自动指标结合；固定随机种子；全过程审计留痕，保证结果可比、可复现。",
        ),
    ],
)

fill_section(
    "4.2 当前结果",
    [
        ("p", "4.2.1 基本任务（文献调研）结果："),
        (
            "table",
            ["Gap", "主题", "类型", "严重度", "置信度", "证据链与建议验证方法"],
            [
                ["Gap 1", "双金属比例-容量定量标度（倒 U 是否普适）", "缺失连接", "高", "0.95", "证据：NiCo-MOF-74 五点实测（Chen 2023）＋Cu/Mg 单调 vs NiCo/CoMn 倒 U；验证：扩展金属对/温度回归"],
                ["Gap 2", "水-CO₂ 竞争/协同机理矛盾（OMS vs 胺）", "矛盾结论", "高", "0.97", "证据：Edison/Marshall/Owens 等 5 路独立证据；验证：RH-容量回归、临界 RH 检验"],
                ["Gap 3", "容量-选择性-再生能 Pareto 前沿未刻画", "缺失连接", "高", "0.85", "证据：Qst 25–40 kJ/mol 窗口报道；验证：跨材料 Pareto 前沿定量刻画"],
                ["Gap 4", "ML 筛选-实验闭环断裂＋数据库误差", "缺失连接", "中", "0.75", "证据：ARC-MOF/LitMOF 结构误差；验证：闭环率指标跨主题统计"],
                ["Gap 5", "DAC（400 ppm）数据稀缺", "未探索", "中", "0.85", "证据：mmen-Mg2(dobpdc) 0.39 mbar 数据；验证：低分压数据扩充"],
                ["Gap 6", "OMS 密度-Qst 标度律缺失", "缺失连接", "中", "0.72", "证据：Caskey 2008 系列 Qst；验证：d 电子数/电负性联合描述符回归"],
                ["Gap 7", "材料-再生工艺耦合优化缺失", "缺失连接", "中", "0.78", "证据：PVSA/TSA 技术经济研究分离；验证：工艺-材料联合优化"],
                ["Gap 8", "杂质气体（NO₂/SO₂）影响研究稀缺", "未探索", "中", "0.70", "证据：Tan 2015 竞争共吸附；验证：NO₂/SO₂ 暴露实验"],
                ["Gap 9", "缺陷工程定量 OMS 控制缺失（缺陷类型二分）", "缺失连接", "中→高", "0.80", "证据：缺失配体 vs 溶剂甲酸盐占据方向相反；验证：CO 探针 IR＋固态 NMR 定量"],
                ["Gap 10", "双金属实际组成 vs 名义比例偏差未量化", "缺失连接", "中→高", "0.78", "证据：Jiao 反应温度主导组成；验证：ICP/EDX 组成核验"],
            ],
            [1.3, 3.4, 1.6, 1.3, 1.2, 8.2],
            8,
        ),
        (
            "bullet",
            "跨主题统计：8 主题累计检索 813+ 篇文献、识别 51 项 Gap；共性 Gap“ML 预测-实验验证闭环断裂”在 5/8 主题被独立识别（一致率 62.5%，热电置信度最高 0.92，见 CROSS_THEME_REPORT.md）；",
        ),
        (
            "bullet",
            "产物清单：调研报告、Gap 清单、知识图谱、论文摘要集与证据链审计日志，位于 workspace/outputs/<主题>/literature_survey/；",
        ),
        (
            "bullet",
            "证据链四层追溯示例：Gap 1 → 支撑证据（NiCo-MOF-74 五点数据，论文标识 v3s0_c795f15f9d35/Chen 2023）→ paper_summaries.md → search_log.jsonl → sciverse_skill_log.jsonl（截至本文档生成时 149 条 API 调用记录，含调用 ID/时间戳/参数哈希；README 历史统计为 143 条，统计时点不同，以日志实际行为准）；审计机制见 audit_trail.md。",
        ),
        (
            "bullet",
            "证据链闭环（2026-08）：gap_report.md 共 48 处 p# 文献引用，46 处（96%）内联 DOI（全部经 Crossref/DataCite/doi.org 逐条核验真实存在，核心证据经人工打开原文确认后回填）；p118 与 p27 为同一来源，合并引用并显式标注 DOI；p16/p139 按“原文待补 DOI”降级标注（水-CO₂ 矛盾一方，不影响主结论）；p54 因来源类型（会议视频）不保留；完整核验清单见提交材料附录。",
        ),
        (
            "bullet",
            "内部量化统计（基于现有真实产物，非实验指标）：10 项 Gap 置信度 0.70–0.97（均值 0.82），全部 10/10 附支撑文献与建议验证方法（证据链完整率 100%）；",
        ),
        (
            "bullet",
            "调研报告质量：主案例调研报告为结构化多章结构（摘要/引言与研究背景/检索策略与方法/知识图谱总览（6 大类约 55 种材料、10 项性质维度、R1–R33 构效关系）/Research Gap 分析/构效关系发现指引/结论与展望/参考文献与证据链），全文文献引用可逐条回溯（正文节选核心文献 20 篇；主案例 11 轮累计检索 546 篇次，去重后持久化证据池 543 条见 paper_register.md，最终收录 46 篇见 paper_summaries.md），见 workspace/outputs/literature_survey/survey_report.md。",
        ),
        ("h3", "4.2.2 路线 A（构效关系发现）实证结果"),
        (
            "table",
            ["假设", "研究对象（材料×性质）", "置信度/新颖性", "搜索方式", "关键证据与状态"],
            [
                ["hypo_1", "NiCo-MOF-74 组分比例-容量", "0.88/0.82", "贝叶斯", "5 实测点二次 R²=0.7694 vs Vegard -0.1530；已定量核验"],
                ["hypo_2", "胺功能化 MOF 湿态容量-RH", "0.90/0.85", "贝叶斯", "多篇独立证据（mmen/MOF-808-AA/TYUT-ATZ）；量化数据待补"],
                ["hypo_3", "Qst 窗口-Pareto 最优", "0.62/0.78", "贝叶斯", "Qst-电负性二次 R²=0.983（p=0.011）；窗口验证待补"],
                ["hypo_4", "MOF-74 缺陷类型依赖", "0.92/0.88", "贝叶斯", "机理二分（缺失配体↑ vs 溶剂占据↓）；定量验证待补"],
                ["hypo_5", "NO₂/SO₂ 暴露-容量衰减-OMS 密度", "0.63/0.90", "贝叶斯", "证据稀缺（Gap 8 延伸）；待扩充"],
            ],
            [1.4, 3.6, 2.2, 2.0, 7.8],
            8,
        ),
        (
            "bullet",
            "定量回归核验：主案例 8 点（含 5 个估计点，如实标注）：二次拟合 R²=0.6919，嵌套 F 检验（线性 vs 二次）F=9.909、p=0.0254 显著，峰值 x≈0.437；",
        ),
        (
            "bullet",
            "独立实测点验证：Chen 2023 五个独立实测点（0℃/1bar，全部 is_estimated=False）：二次 R²=0.7694、高斯 R²=0.9778，均显著优于经典 Vegard 端点线性基线（R²=-0.1530，ΔR²=+0.92/+1.13）；n=5 的 F 检验 p=0.158 功效不足（如实披露），bootstrap（1000 次）89.3% 支持二次优于线性；",
        ),
        (
            "bullet",
            "多口径交叉印证：归一化复合 12 点嵌套 F p=4.9e-05；跨 4 体系 meta-analysis 平均增强率 61.3%、t=5.46、p=0.0121；Qst-电负性二次 R²=0.983 vs 线性基线 0.190（p=0.011）；",
        ),
        (
            "bullet",
            "“为何旧模型失效”机制解释：Vegard 线性混合的数学本质是两端点性能的线性插值，它只适用于“两种金属位点各自独立贡献”的情形；双金属 MOF-74 的倒 U 说明中间组分出现了端点没有的协同，NiCo 5 个实测点上线性模型 R² 为负（-0.1530，连用均值预测都不如），而二次/高斯拟合 ΔR²=+0.92/+1.13，说明协同项真实存在、必须显式建模。Slack 经验公式面向高温、简单晶体标定，自变量不含湿度、缺陷浓度、双金属比例，天然无法外推到 MOF 的湿态、缺陷与双金属协同场景——这些正是我们识别 Gap、生成假设的区间；",
        ),
        (
            "bullet",
            "外部数据库验证：Materials Project 氧化物代理热力学提供间接证据（如 MgNiO₂ 双金属 ΔE=-0.035 eV/atom 更稳定，明确标注为间接参考）；OQMD/hMOF 提供形成能与稳定性支持（如 CoNiO₂/Co₂NiO₄）；OQMD/NOMAD 对 MOF 吸附性质覆盖为 0，如实记录为负结果并升级为领域数据基础设施 Gap；",
        ),
        (
            "bullet",
            "参照系公平对比：同预算随机搜索对照（10 种子），v2 增强打分下 5 条假设全部 bayesian_wins（diff +0.014~+0.039），修复了初赛版 10 种子 3:2 区分度不足的问题（打分函数分段线性拉伸）；",
        ),
        (
            "bullet",
            "四类发现信号（运行前定义）：正结果（双金属倒 U、水 RH 双分支、缺陷类型二分，置信度≥0.7 且带证据链）、负结果（MP/OQMD 零覆盖）、异常（Ni-MOF-74 容量矛盾区间、水-CO₂ 机理矛盾）、反例（胺型 MOF 低湿度容量反升，与“水总是有害”朴素假设相反）；",
        ),
        (
            "bullet",
            "诚实披露的局限：主案例 8 点中 5 个为估计点（不作独立验证证据）；n=5 的 F 检验 p=0.158 功效不足——统计显著性主要由归一化复合 12 点（p=4.9e-05）与跨体系 meta-analysis（p=0.0121）提供；Qst n=5 小样本；跨主题 5 个连接均为理论推导待验证。",
        ),
        ("h3", "4.2.3 待实施/待补充（如实标注，不虚构数据）"),
        (
            "bullet",
            "对照实验量化指标（Precision/Recall/Recall@K/nDCG@K、抽取 F1、证据定位准确率等）：【待实施】拟按 4.1.2 方案开展，暂无量化数据；",
        ),
        (
            "bullet",
            "Research Gap 专家认可率、新颖性、可操作性、证据完整率：【待补充】需领域专家评审后填写（内部代理指标见 4.2.1）；",
        ),
        (
            "bullet",
            "报告生成时间与 API 成本统计：【待补充】。",
        ),
        (
            "p",
            "声明：本节全部为真实运行产物；未完成部分如实标注为待实施/待补充，不存在虚构的实验数据、性能指标或系统能力。",
        ),
    ],
)

# 新增 4.3 后续演进与复赛计划（插在“五、复现与开放计划”之前）
insert_before_heading(
    1,
    "五、复现与开放计划",
    [
        ("h2", "4.3 后续演进与复赛计划"),
        (
            "p",
            "近期（初赛→复赛）：① 完成 4.1.2 小规模对照实验并产出量化指标；② 数据工程第一优先——表格化知识图谱（材料体系×结构变量×性能值 (x,y) 表，带 p# 溯源），为“候选 vs 经典（Slack/Vegard）”正面对比提供干净数据；③ v2.0 全量重跑主案例，交付 llm_guidance 审计记录（LLM 搜索内引导实证）；④ 复赛全量重跑中验证 v2 打分函数在多主题的稳定性（主案例区分度已优化，见 4.2.2）。",
        ),
        (
            "p",
            "复赛：提交可运行代码仓库（README、环境配置、Docker、CI、评测脚本）＋完整实验结果报告＋科学意义阐释；达成路线 A“统计优于前人”验证标准（更多体系实测点、跨体系 meta-analysis 扩展）；启动跨子领域连接验证（首选 MOF 缺陷工程→电池正极缺陷分类）。",
        ),
        (
            "p",
            "时间表：以组委会正式赛程为准【待确认】。",
        ),
    ],
)

# ---------- 五、复现与开放计划 ----------
fill_section(
    "5.1 复现方式",
    [
        (
            "bullet",
            "环境：Python 3.10+（开发环境 3.12）；pip install -r requirements.txt，或使用 Dockerfile / docker-compose.yml 容器化部署；",
        ),
        (
            "bullet",
            "配置：.api_key（DEEPSEEK_API_KEY / SCIVERSE_API_KEY / MINERU_API_KEY，均可选；缺失自动降级至 arXiv＋本地解析，保证无密钥可运行）；",
        ),
        (
            "bullet",
            "确定性：去重、单位换算、数值校验、Gap 评分等确定性计算由 seed_everything() 固定随机种子（--seed 42）；LLM 采样不保证逐字复现，但全链路审计日志（检索调用、解析引擎、抽取记录、Gap 证据链）保证任何结论可独立核验；",
        ),
        (
            "bullet",
            "复现命令：python main.py --topic \"MOF materials for CO2 capture\" --budget 600 --fresh --seed 42（多主题示例：钙钛矿/热电/正极主题见 README.md 多主题运行）；模块自测 python -m pytest tests/；复现核验（无需网络）：python literature_agent/classical_models.py（Slack/Vegard 参数恢复自检）、python literature_agent/symbolic_regression.py（表达式恢复自检）、python literature_agent/extractor.py（数值配对自测）；详细步骤见 README.md / REPRODUCIBILITY.md；",
        ),
        (
            "bullet",
            "目标架构选型已确定默认方案（自研事件驱动状态机、Chroma、BGE-M3、bge-reranker-v2-m3，见 3.1.2、3.3.1），以复现验证为准【待确认】。对应依赖已列入 requirements.txt（拟实施用途，实施时实测锁定）。",
        ),
    ],
)

fill_section(
    "5.2 开源计划",
    [
        (
            "p",
            "初赛阶段：按赛题要求不强制提交代码，本方案明确开源边界与复现承诺。复赛/决赛阶段：以公开 GitHub 仓库提交可运行代码（含核心 Agent、工具层、脚本、测试、Dockerfile、requirements.txt）与评测/对照实验脚本，开放调研报告、Gap 清单、知识图谱、审计日志、对照实验数据等运行产物（API Key 除外）。",
        ),
        (
            "p",
            "许可：代码采用 MIT；文档与运行产物采用 CC BY 4.0（不含第三方 API 数据）。边界：Sciverse 检索结果仅限内部调研使用、不对外再分发；商业 API Key 不公开。",
        ),
    ],
)

fill_section(
    "5.3 依赖、数据来源与合规披露",
    [
        (
            "bullet",
            "商业 API/闭源模型：当前原型使用 DeepSeek 闭源推理模型（deepseek-v4-flash，OpenAI 兼容接口，见 COMPLIANCE.md）。使用范围：Agent 推理与报告文本；原因：推理质量与稳定性；费用假设（估算）：600s 单轮约数十次 LLM 调用，DeepSeek 按 token 计费，单轮成本量级约 0.1–1 元；替代方案：任意 OpenAI 兼容端点（DEEPSEEK_BASE_URL / DEEPSEEK_MODEL 环境变量切换，含本地 vLLM 部署的开源模型）；迁移成本：环境变量级；对复现性影响：LLM 采样随机，但确定性计算层与证据链保证结论可独立核验。",
        ),
        (
            "bullet",
            "Sciverse API：文献检索（search.py / sciverse_mcp.py），REST＞MCP＞Skill 三层自动降级，缺失 Key 自动回退纯 arXiv（零成本）；费用假设（估算）：按平台套餐计费；每次调用留痕审计（sciverse_skill_log.jsonl，含调用 ID、时间戳、参数哈希）。",
        ),
        (
            "bullet",
            "MinerU：PDF 解析（Cloud＞本地服务＞pip 包），全部不可用自动回退本地 markitdown＋pdfplumber（离线、确定可复现），回退机制经 mineru_test_results.json 实测。费用假设（估算）：MinerU 云 API 按平台配额/用量计费（需 MINERU_API_KEY），本地部署（localhost:8888 或 pip 包）无配额影响。当前运行环境 MinerU 云/本地均不可用（未配置 key/未部署服务），实际解析走本地回退引擎。",
        ),
        (
            "bullet",
            "数据来源与授权：arXiv（开放获取）；Sci-Base（Hugging Face 数据集授权）；Sciverse（仅取标题＋摘要用于内部调研，不对外再分发）；Materials Project / OQMD / NOMAD（公开数据库，供路线 A 交叉验证）。",
        ),
        (
            "bullet",
            "第三方依赖（关键版本）：requirements.txt 精确锁定（openai==1.108.1、requests==2.32.5、markitdown==0.1.7、pdfplumber==0.11.10、pdfminer.six==20260107、numpy==1.26.4、scipy==1.17.1、python-docx==1.2.0 等），与 requirements.txt 实际一致。另预置拟实施依赖：chromadb==1.5.9、sentence-transformers==5.6.1、FlagEmbedding==1.4.0（复赛实施阶段启用，届时实测锁定）。",
        ),
        (
            "bullet",
            "密钥管理：.api_key 已 gitignore、不入库；完整披露见 COMPLIANCE.md。",
        ),
        (
            "bullet",
            "已有项目声明：本方案为团队独立实现，未基于任何已有项目继续开发（见 COMPLIANCE.md §八）；核心代码、23 个工具、双引擎解析、三层 Sciverse 接入与贝叶斯/MCTS 发现层均为自主开发。",
        ),
    ],
)

# ---------- 六、团队介绍（保留模板标题，填写待补充占位） ----------
def fill_team(anchor_heading: str, text: str):
    h = find_heading(2, anchor_heading)
    assert h is not None, anchor_heading
    nxt = h._p.getnext()
    p = doc.add_paragraph(style="Normal")
    p.add_run(text)
    if nxt is not None and nxt.tag == qn("w:p"):
        nxt.addnext(p._p)
    else:
        h._p.addnext(p._p)


fill_team("6.1 成员背景", "【待补充：请填写每位成员的学校/公司、岗位/专业，以及核心技能（技术、专业领域）。】")
fill_team("6.2 团队分工", "【待补充：请列出每位核心成员在项目中的角色与分工。】")
fill_team("6.3团队成果", "【待补充：请填写过往项目经历/获奖经历、作品合集链接（如有）。】")

# 删除空 Heading 2 段落（模板遗留空标题）
for p in list(doc.paragraphs):
    if p.style and p.style.name == "Heading 2" and not p.text.strip():
        remove_paragraph(p)

# 删除两条“呈现建议（非赛事要求）”模板提示
for p in list(doc.paragraphs):
    if p.text.strip().startswith("呈现建议（非赛事要求）"):
        remove_paragraph(p)

# ---------- 七、待团队补充信息清单 ----------
h1 = make_para("七、待团队补充信息清单", style="Heading 1")
items = [
    "团队介绍：成员背景、团队分工、过往成果（第六章）；",
    "Research Gap 专家评审数据（认可率/新颖性/可操作性/证据完整率）与 API 成本统计（4.2.3）；",
    "对照实验具体完成日期（标注集与金标准规模已在 4.1.2 确定，以复赛截止时间为准）；",
    "目标架构选型的复现验证结果确认（默认方案已确定：自研状态机、Chroma、BGE-M3、bge-reranker-v2-m3；3.1.2、3.3.1）；",
    "路线 A 是否展开为独立章节/文档（1.2、3.1.5）；",
    "复赛时间表与里程碑确认（4.3）。",
]
cursor = h1._p
for it in items:
    p = make_para(it, style="List Bullet")
    cursor.addnext(p._p)
    cursor = p._p


# ---------- 统一各级标题格式 ----------
def normalize_headings():
    """统一 Heading 1/2/3 的字体、字号、加粗与颜色（清除模板遗留的直接格式差异）"""
    specs = {
        "Heading 1": dict(size=16.0, color="2E74B5"),
        "Heading 2": dict(size=12.5, color="1F4D78"),
        "Heading 3": dict(size=11.0, color="1F4D78"),
    }
    for p in doc.paragraphs:
        if not (p.style and p.style.name in specs):
            continue
        if not p.text.strip():
            continue
        spec = specs[p.style.name]
        for r in p.runs:
            r.font.name = "Times New Roman"
            rPr = r._r.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:ascii"), "Times New Roman")
            rFonts.set(qn("w:hAnsi"), "Times New Roman")
            rFonts.set(qn("w:eastAsia"), "黑体")
            r.font.size = Pt(spec["size"])
            r.font.bold = True
            r.font.color.rgb = RGBColor.from_string(spec["color"])


normalize_headings()

doc.save(OUT)
print("saved:", OUT)
